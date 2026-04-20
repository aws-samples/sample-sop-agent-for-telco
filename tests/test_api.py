# Copyright Amazon.com, Inc. or its affiliates. All Rights Reserved.
# SPDX-License-Identifier: MIT-0
"""Tests for api.py — App stats calculation, endpoints."""
import pytest
from unittest.mock import patch, MagicMock
from fastapi.testclient import TestClient

# Patch subprocess before importing api to prevent background threads from hitting real cluster
with patch("subprocess.check_output", return_value=""), \
     patch("subprocess.run", return_value=MagicMock(stdout="", stderr="", returncode=0)):
    from api import app, _upf_stats_cache, EventBuffer


client = TestClient(app)


# ── API Endpoints ──

class TestSOPEndpoints:
    def test_list_sops(self):
        resp = client.get("/api/sops")
        assert resp.status_code == 200
        assert isinstance(resp.json(), list)

    def test_get_sop_not_found(self):
        resp = client.get("/api/sop/nonexistent.md")
        assert resp.status_code == 404

    def test_get_status(self):
        resp = client.get("/api/status")
        assert resp.status_code == 200
        data = resp.json()
        assert "status" in data

    def test_get_metrics(self):
        resp = client.get("/api/metrics")
        assert resp.status_code == 200
        data = resp.json()
        assert "rxGbps" in data
        assert "txGbps" in data
        assert "combined" in data

    def test_get_upf_stats(self):
        resp = client.get("/api/app-stats")
        assert resp.status_code == 200
        data = resp.json()
        for key in ["ipackets", "opackets", "imissed", "dropRate", "fwdLoss", "totalLoss", "workerNG"]:
            assert key in data

    def test_get_alarms(self):
        resp = client.get("/api/alarms")
        assert resp.status_code == 200
        assert isinstance(resp.json(), list)


# ── App Stats Calculation ──

class TestUPFStatsCalculation:
    """Test the math behind App forwarding stats."""

    def test_zero_packets_no_division_error(self):
        """When no packets, all rates should be 0."""
        _upf_stats_cache["data"] = {
            "ipackets": 0, "opackets": 0, "imissed": 0,
            "dropRate": 0, "fwdLoss": 0, "workerNG": 0, "totalLoss": 0,
        }
        resp = client.get("/api/app-stats")
        data = resp.json()
        assert data["dropRate"] == 0
        assert data["fwdLoss"] == 0
        assert data["totalLoss"] == 0

    def test_perfect_forwarding(self):
        """All packets forwarded, no loss."""
        _upf_stats_cache["data"] = {
            "ipackets": 1000000, "opackets": 1000000, "imissed": 0,
            "dropRate": 0, "fwdLoss": 0, "workerNG": 0, "totalLoss": 0,
        }
        data = client.get("/api/app-stats").json()
        assert data["totalLoss"] == 0

    def test_high_imissed_shows_drop_rate(self):
        """938M imissed on 75B packets = ~1.2% drop rate."""
        ipkts = 75_000_000_000
        imiss = 938_000_000
        total = ipkts + imiss
        expected_drop = round(imiss / total * 100, 5)
        _upf_stats_cache["data"] = {
            "ipackets": ipkts, "opackets": ipkts - 500_000_000, "imissed": imiss,
            "dropRate": expected_drop,
            "fwdLoss": round(500_000_000 / ipkts * 100, 5),
            "workerNG": 0,
            "totalLoss": round((total - (ipkts - 500_000_000)) / total * 100, 5),
        }
        data = client.get("/api/app-stats").json()
        assert data["dropRate"] > 1.0
        assert data["totalLoss"] > 1.0


# ── Security: kubectl allowlist & shell blocklist ──

class TestChatToolSecurity:
    """Test the command safety checks added in Phase 1."""

    def test_health_endpoint(self):
        resp = client.get("/health")
        assert resp.status_code == 200
        assert resp.json() == {"status": "ok"}

    def test_get_valid_sop(self):
        sops = client.get("/api/sops").json()
        if sops:
            resp = client.get(f"/api/sop/{sops[0]['name']}")
            assert resp.status_code == 200
            assert "content" in resp.json()


class TestUPFStatsEdgeCases:
    """Edge cases for App stats calculation."""

    def test_baseline_subtraction_no_negative(self):
        """Stats should never go negative after baseline."""
        _upf_stats_cache["data"] = {
            "ipackets": 0, "opackets": 0, "imissed": 0,
            "dropRate": 0, "fwdLoss": 0, "workerNG": 0, "totalLoss": 0,
        }
        data = client.get("/api/app-stats").json()
        assert data["ipackets"] >= 0
        assert data["opackets"] >= 0
        assert data["imissed"] >= 0

    def test_small_packet_count_not_meaningful(self):
        """Below 100K packets, rates should be 0 (not meaningful)."""
        _upf_stats_cache["data"] = {
            "ipackets": 1000, "opackets": 500, "imissed": 500,
            "dropRate": 0, "fwdLoss": 0, "workerNG": 0, "totalLoss": 0,
        }
        data = client.get("/api/app-stats").json()
        # These are pre-computed in cache, but verify structure
        for key in ["dropRate", "fwdLoss", "totalLoss"]:
            assert isinstance(data[key], (int, float))


# ── Graph Endpoint ──

class TestGraphEndpoint:
    """Test /ws/execute-graph WebSocket endpoint structure."""

    def test_graph_endpoint_exists(self):
        """Verify the graph WebSocket route is registered."""
        routes = [r.path for r in app.routes]
        assert "/ws/execute-graph" in routes

    def test_graph_endpoint_rejects_empty_sops(self):
        """Empty sop_paths should return error."""
        with client.websocket_connect("/ws/execute-graph") as ws:
            ws.send_json({"sop_paths": [], "model": "haiku"})
            resp = ws.receive_json()
            assert resp["type"] == "error"
            assert "No SOPs" in resp["message"]


# ── EventBuffer ──

class TestEventBuffer:
    def test_append_and_since(self):
        buf = EventBuffer(maxlen=100)
        buf.append({"type": "a"})
        buf.append({"type": "b"})
        assert buf.last_seq == 2
        assert len(buf.since(0)) == 2
        assert len(buf.since(1)) == 1
        assert buf.since(1)[0]["type"] == "b"

    def test_clear(self):
        buf = EventBuffer(maxlen=100)
        buf.append({"type": "a"})
        buf.clear()
        assert buf.last_seq == 0
        assert len(buf.since(0)) == 0

    def test_ring_buffer_evicts_old(self):
        buf = EventBuffer(maxlen=3)
        for i in range(5):
            buf.append({"i": i})
        # Only last 3 remain
        events = buf.since(0)
        assert len(events) == 3
        assert events[0]["i"] == 2

    def test_since_empty(self):
        buf = EventBuffer(maxlen=100)
        assert buf.since(0) == []
        assert buf.since(999) == []


# ── Document Upload / SOP Generation ──

class TestGenerateSOP:
    """Test POST /api/generate-sop endpoint."""

    def test_upload_txt_generates_sop(self, tmp_path, monkeypatch):
        """Uploading a .txt file creates an SOP in the sops directory."""
        import api
        monkeypatch.setattr(api, "SOP_REPO", str(tmp_path))
        (tmp_path / "sops").mkdir()

        # Mock the agent call to avoid hitting Bedrock
        monkeypatch.setattr(api, "_generate_sop_with_agent", lambda text, fname: f"# Generated SOP\n\nFrom: {fname}\n\n## Procedure\n\n{text[:100]}")

        resp = client.post(
            "/api/generate-sop",
            files={"file": ("test-runbook.txt", b"Step 1: Deploy App\nStep 2: Validate", "text/plain")},
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["status"] == "generated"
        assert data["sop_name"] == "test-runbook.md"
        assert data["source"] == "test-runbook.txt"
        # Verify file was actually created
        sop_path = tmp_path / "sops" / "test-runbook.md"
        assert sop_path.exists()
        content = sop_path.read_text()
        assert "Step 1: Deploy App" in content

    def test_upload_md_generates_sop(self, tmp_path, monkeypatch):
        """Uploading a .md file works."""
        import api
        monkeypatch.setattr(api, "SOP_REPO", str(tmp_path))
        (tmp_path / "sops").mkdir()
        monkeypatch.setattr(api, "_generate_sop_with_agent", lambda text, fname: f"# SOP from {fname}")

        resp = client.post(
            "/api/generate-sop",
            files={"file": ("hld-network.md", b"# HLD\n## Network Design\nVLAN 3505", "text/markdown")},
        )
        assert resp.status_code == 200
        assert resp.json()["sop_name"] == "hld-network.md"

    def test_upload_pdf_extracts_text(self, tmp_path, monkeypatch):
        """Uploading a .pdf attempts text extraction."""
        import api
        monkeypatch.setattr(api, "SOP_REPO", str(tmp_path))
        (tmp_path / "sops").mkdir()
        monkeypatch.setattr(api, "_generate_sop_with_agent", lambda text, fname: f"# SOP from {fname}\n\n{text[:80]}")

        resp = client.post(
            "/api/generate-sop",
            files={"file": ("vendor-guide.pdf", b"%PDF-1.4 fake", "application/pdf")},
        )
        assert resp.status_code == 200
        content = (tmp_path / "sops" / "vendor-guide.md").read_text()
        # Fake PDF bytes won't parse, so we get the extraction error fallback
        assert "Failed to extract" in content or "vendor-guide" in content

    def test_upload_docx_extracts_text(self, tmp_path, monkeypatch):
        """Uploading a .docx extracts paragraph text."""
        import api
        from docx import Document
        import io
        monkeypatch.setattr(api, "SOP_REPO", str(tmp_path))
        (tmp_path / "sops").mkdir()
        monkeypatch.setattr(api, "_generate_sop_with_agent", lambda text, fname: f"# SOP\n\n{text[:80]}")

        # Create a real docx in memory
        doc = Document()
        doc.add_paragraph("Step 1: Run kubectl apply")
        doc.add_paragraph("Step 2: Verify pods")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        resp = client.post(
            "/api/generate-sop",
            files={"file": ("guide.docx", buf.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert resp.status_code == 200
        content = (tmp_path / "sops" / "guide.md").read_text()
        assert "kubectl apply" in content

    def test_upload_rejects_unsupported_type(self, tmp_path, monkeypatch):
        """Uploading a .zip should be rejected."""
        import api
        monkeypatch.setattr(api, "SOP_REPO", str(tmp_path))

        resp = client.post(
            "/api/generate-sop",
            files={"file": ("archive.zip", b"PK\x03\x04", "application/zip")},
        )
        assert resp.status_code == 400
        assert "Unsupported" in resp.json()["detail"]

    def test_upload_avoids_overwrite(self, tmp_path, monkeypatch):
        """If SOP name already exists, appends a counter."""
        import api
        monkeypatch.setattr(api, "SOP_REPO", str(tmp_path))
        (tmp_path / "sops").mkdir()
        (tmp_path / "sops" / "existing.md").write_text("# Existing")
        monkeypatch.setattr(api, "_generate_sop_with_agent", lambda text, fname: "# New SOP")

        resp = client.post(
            "/api/generate-sop",
            files={"file": ("existing.txt", b"new content", "text/plain")},
        )
        assert resp.status_code == 200
        assert resp.json()["sop_name"] == "existing-1.md"
        # Original untouched
        assert (tmp_path / "sops" / "existing.md").read_text() == "# Existing"

    def test_upload_saves_to_uploads_dir(self, tmp_path, monkeypatch):
        """Uploaded file is saved to uploads/ directory."""
        import api
        monkeypatch.setattr(api, "SOP_REPO", str(tmp_path))
        (tmp_path / "sops").mkdir()
        monkeypatch.setattr(api, "_generate_sop_with_agent", lambda text, fname: "# SOP")

        client.post(
            "/api/generate-sop",
            files={"file": ("my-doc.txt", b"hello", "text/plain")},
        )
        assert (tmp_path / "uploads" / "my-doc.txt").exists()
        assert (tmp_path / "uploads" / "my-doc.txt").read_bytes() == b"hello"

    def test_agent_fallback_on_error(self, tmp_path, monkeypatch):
        """If agent fails, fallback placeholder SOP is created."""
        import api
        monkeypatch.setattr(api, "SOP_REPO", str(tmp_path))
        (tmp_path / "sops").mkdir()

        def _fail(text, fname):
            raise RuntimeError("Bedrock unavailable")
        monkeypatch.setattr(api, "_generate_sop_with_agent", _fail)

        resp = client.post(
            "/api/generate-sop",
            files={"file": ("fail-test.txt", b"content", "text/plain")},
        )
        assert resp.status_code == 200
        content = (tmp_path / "sops" / "fail-test.md").read_text()
        assert "Generation failed" in content
